#!/usr/bin/env python3
"""Build the Spanish KDP interior (6x9") DOCX from the translated markdown files."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BOOK="/home/user/sophie-ui/book/es"
PARTS=["01-front-and-intro","02-part1","03-part2","04-part3","05-part4","06-backmatter"]

def add_field(p,instr):
    r=p.add_run()
    for tag,attr,val,txt in [("w:fldChar","w:fldCharType","begin",None),(None,None,None,None)]:
        pass
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    t=OxmlElement('w:t'); t.text=""
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    for e in (f1,it,f2,t,f3): r._r.append(e)

def add_inline(p,text):
    for tok in re.split(r'(\*\*.*?\*\*|\*.*?\*)',text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'): p.add_run(tok[2:-2]).bold=True
        elif tok.startswith('*') and tok.endswith('*'): p.add_run(tok[1:-1]).italic=True
        else: p.add_run(tok)

doc=Document()
n=doc.styles['Normal']; n.font.name='Georgia'; n.font.size=Pt(11)
pf=n.paragraph_format; pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; pf.line_spacing=1.15; pf.space_after=Pt(0); pf.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
body=doc.styles.add_style('Body',1); body.base_style=doc.styles['Normal']; body.font.name='Georgia'; body.font.size=Pt(11)
body.paragraph_format.first_line_indent=Inches(0.22); body.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; body.paragraph_format.line_spacing=1.15
for name,size,before,after in [('PartHead',24,36,18),('ChapHead',17,30,16)]:
    st=doc.styles.add_style(name,1); st.base_style=doc.styles['Normal']; st.font.name='Georgia'; st.font.size=Pt(size); st.font.bold=True
    st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    st.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
sub=doc.styles.add_style('SubHead',1); sub.base_style=doc.styles['Normal']; sub.font.name='Georgia'; sub.font.size=Pt(12.5); sub.font.bold=True
sub.paragraph_format.space_before=Pt(12); sub.paragraph_format.space_after=Pt(4); sub.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

sec=doc.sections[0]; sec.page_width=Inches(6); sec.page_height=Inches(9)
sec.top_margin=Inches(0.75); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.875); sec.right_margin=Inches(0.5); sec.gutter=Inches(0.13)
doc.settings.element.append(OxmlElement('w:mirrorMargins'))
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_field(fp,"PAGE")

def new_page(): doc.add_page_break()
def toc_h(level,text):
    p=doc.add_paragraph(style=('ChapHead' if level==2 else 'PartHead')); add_inline(p,text)
    ol=OxmlElement('w:outlineLvl'); ol.set(qn('w:val'),str(level-1)); p._p.get_or_add_pPr().append(ol)

# TITLE PAGE
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("MORMONES"); r.bold=True; r.font.size=Pt(30)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Los Santos de los Últimos Días, Explicados"); r.bold=True; r.font.size=Pt(18)
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Una guía clara y ecuánime sobre las creencias,\nla historia y la vida cotidiana de los mormones"); r.italic=True; r.font.size=Pt(13)
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("SOPHIE E T"); r.font.size=Pt(15)
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Una guía educativa independiente"); r.italic=True; r.font.size=Pt(11)
new_page()

# COPYRIGHT
cp=[
 "Copyright © 2026. Todos los derechos reservados.",
 "",
 "Ninguna parte de este libro puede reproducirse, distribuirse ni transmitirse de ninguna forma ni por ningún medio, incluidos el fotocopiado, la grabación u otros métodos electrónicos o mecánicos, sin el permiso previo por escrito del editor, salvo en el caso de breves citas incluidas en reseñas críticas y ciertos otros usos no comerciales permitidos por la ley de derechos de autor.",
 "",
 "Nota sobre la independencia. Esta es una obra independiente de divulgación y periodismo. NO está publicada, patrocinada, autorizada ni respaldada por La Iglesia de Jesucristo de los Santos de los Últimos Días, ninguna de sus organizaciones afiliadas, ni ninguna otra iglesia, confesión o grupo mencionado en estas páginas. Todas las marcas y nombres son propiedad de sus respectivos titulares y se usan aquí solo con fines de identificación y comentario. Las interpretaciones y cualquier error son responsabilidad de la autora.",
 "",
 "Nota sobre las fuentes. Este libro es una síntesis de información pública ampliamente disponible: las declaraciones oficiales de la Iglesia y sus ensayos de Temas del Evangelio, obras de referencia estándar, periodismo de medios generales y el trabajo de historiadores de la religión estadounidense. No cita extensamente ninguna escritura protegida por derechos de autor ni reproduce texto de propiedad exclusiva. Cuando los estudiosos o los creyentes discrepan de verdad, el libro lo dice en vez de tomar partido.",
 "",
 "Primera edición — 2026",
]
for _ in range(2): doc.add_paragraph()
for line in cp:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.line_spacing=1.15
    if line:
        m=re.match(r'^(Nota sobre [^.]+\.)(.*)',line)
        if m: p.add_run(m.group(1)).bold=True; p.add_run(m.group(2))
        else: p.add_run(line)
    for rr in p.runs: rr.font.size=Pt(9.5)
new_page()

# TOC
p=doc.add_paragraph(style='ChapHead'); p.add_run("Contenido")
p2=doc.add_paragraph(); add_field(p2,'TOC \\o "1-2" \\h \\z \\u')
doc.settings.element.append(OxmlElement('w:updateFields')); doc.settings.element[-1].set(qn('w:val'),'true')
new_page()

# CONTENT
chunks=[]
for i,name in enumerate(PARTS):
    txt=open(f"{BOOK}/{name}.md",encoding="utf-8").read()
    if i==0:
        m=re.search(r'(?m)^## ',txt)  # start at first H2 (Nota sobre los nombres)
        if m: txt=txt[m.start():]
    chunks.append(txt)
lines="\n\n".join(chunks).split("\n")
i=0
while i<len(lines):
    line=lines[i].rstrip(); s=line.strip()
    if s=="---" or s=="": i+=1; continue
    if re.match(r'^# ',line): new_page(); toc_h(1,s[2:].strip()); i+=1; continue
    if re.match(r'^## ',line): new_page(); toc_h(2,s[3:].strip()); i+=1; continue
    if re.match(r'^### ',line):
        p=doc.add_paragraph(style='SubHead'); add_inline(p,s[4:].strip()); i+=1; continue
    if s.startswith(">"):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.4); p.paragraph_format.right_indent=Inches(0.3)
        p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6); add_inline(p,s.lstrip(">").strip())
        for r in p.runs: r.italic=True
        i+=1; continue
    if s.startswith("- "):
        p=doc.add_paragraph(style='List Bullet'); add_inline(p,s[2:].strip()); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT; i+=1; continue
    m=re.match(r'^\d+\.\s+(.*)',s)
    if m:
        p=doc.add_paragraph(style='List Number'); add_inline(p,m.group(1)); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT; i+=1; continue
    p=doc.add_paragraph(style='Body'); add_inline(p,s); i+=1

doc.save(f"{BOOK}/KDP-interior-6x9.docx"); print("saved es/KDP-interior-6x9.docx")
