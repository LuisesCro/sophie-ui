#!/usr/bin/env python3
"""Build a KDP print-ready interior (6x9") DOCX from the markdown chapter files."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BOOK_DIR = "/home/user/sophie-ui/book"
PARTS = ["01-front-and-intro", "02-part1", "03-part2", "04-part3", "05-part4", "06-backmatter"]

# ---------- helpers ----------
def add_field(paragraph, instr):
    """Insert a Word field (used for TOC and page numbers)."""
    r = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = instr
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = ""
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1); r._r.append(instrText); r._r.append(fldChar2); r._r.append(t); r._r.append(fldChar3)

def set_updatefields(doc):
    settings = doc.settings.element
    el = OxmlElement('w:updateFields'); el.set(qn('w:val'), 'true')
    settings.append(el)

def add_inline(paragraph, text):
    """Add runs parsing **bold** and *italic*."""
    # split keeping delimiters
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2]); run.bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            run = paragraph.add_run(tok[1:-1]); run.italic = True
        else:
            paragraph.add_run(tok)

# ---------- document + styles ----------
doc = Document()

# Base style
normal = doc.styles['Normal']
normal.font.name = 'Georgia'
normal.font.size = Pt(11)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Body style (first-line indent)
body = doc.styles.add_style('Body', 1)
body.base_style = doc.styles['Normal']
body.font.name = 'Georgia'; body.font.size = Pt(11)
body.paragraph_format.first_line_indent = Inches(0.22)
body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
body.paragraph_format.line_spacing = 1.15

# Headings
for name, size, before, after in [('PartHead', 24, 36, 18), ('ChapHead', 17, 30, 16)]:
    st = doc.styles.add_style(name, 1)
    st.base_style = doc.styles['Normal']
    st.font.name = 'Georgia'; st.font.size = Pt(size); st.font.bold = True
    st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_before = Pt(before); st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    # map to builtin heading for TOC
    st.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

sub = doc.styles.add_style('SubHead', 1)
sub.base_style = doc.styles['Normal']
sub.font.name = 'Georgia'; sub.font.size = Pt(12.5); sub.font.bold = True
sub.paragraph_format.space_before = Pt(12); sub.paragraph_format.space_after = Pt(4)
sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ---------- page setup: 6x9 mirror margins ----------
sec = doc.sections[0]
sec.page_width = Inches(6); sec.page_height = Inches(9)
sec.top_margin = Inches(0.75); sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.875); sec.right_margin = Inches(0.5)
sec.gutter = Inches(0.13)
# mirror margins
pgMar = sec._sectPr.find(qn('w:pgMar'))
sectPr = sec._sectPr
mirror = OxmlElement('w:mirrorMargins');
# mirrorMargins goes in settings, not sectPr; set via settings
mm = OxmlElement('w:mirrorMargins')
doc.settings.element.append(mm)

# page number in footer (centered)
footer = sec.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.style = doc.styles['Normal']
add_field(fp, "PAGE")

def new_page():
    doc.add_page_break()

def toc_h(level, text):
    """Add a paragraph that is a real heading (so TOC/Kindle nav picks it up)."""
    style = 'ChapHead' if level == 2 else 'PartHead'
    p = doc.add_paragraph(style=style)
    add_inline(p, text)
    # tag with outline level for TOC
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl'); ol.set(qn('w:val'), str(level-1))
    pPr.append(ol)
    return p

# ---------- TITLE PAGE ----------
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MORMONS"); r.bold = True; r.font.size = Pt(30); r.font.name='Georgia'
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The Latter-day Saints, Explained"); r.bold = True; r.font.size = Pt(20)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A Clear, Fair Guide to Mormon Beliefs, History,\nand Everyday Life"); r.italic = True; r.font.size = Pt(13)
for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("An independent educational guide"); r.font.size = Pt(11); r.italic=True
new_page()

# ---------- COPYRIGHT PAGE ----------
cp_lines = [
    "Copyright © 2026. All rights reserved.",
    "",
    "No part of this book may be reproduced, distributed, or transmitted in any form or by any means, including photocopying, recording, or other electronic or mechanical methods, without the prior written permission of the publisher, except in the case of brief quotations embodied in critical reviews and certain other noncommercial uses permitted by copyright law.",
    "",
    "A note on independence. This is an independent work of general education and journalism. It is not published, sponsored, authorized, endorsed by, or affiliated with The Church of Jesus Christ of Latter-day Saints, any of its affiliated organizations, or any other church, denomination, or group mentioned in these pages. All trademarks and names are the property of their respective owners and are used here only for identification and commentary. The interpretations and any errors are the author's own.",
    "",
    "A note on sources. This book is a synthesis of widely available public information: the church's own published statements and Gospel Topics essays, standard reference works, mainstream journalism, and the work of historians of American religion. It quotes no copyrighted scripture at length and reproduces no proprietary text. Where scholars or believers genuinely disagree, the book says so rather than choosing sides.",
    "",
    "First Edition — 2026",
]
for _ in range(2):
    doc.add_paragraph()
for line in cp_lines:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    if line:
        # bold the lead labels
        m = re.match(r'^(A note on \w+\.)(.*)', line)
        if m:
            r=p.add_run(m.group(1)); r.bold=True; p.add_run(m.group(2))
        else:
            p.add_run(line)
    r0 = p.runs[0] if p.runs else None
    for rr in p.runs: rr.font.size = Pt(9.5)
new_page()

# ---------- TABLE OF CONTENTS ----------
p = doc.add_paragraph(style='ChapHead'); p.add_run("Contents")
p2 = doc.add_paragraph()
add_field(p2, 'TOC \\o "1-2" \\h \\z \\u')
set_updatefields(doc)
new_page()

# ---------- PARSE CONTENT ----------
def read_content():
    chunks = []
    for i, name in enumerate(PARTS):
        with open(f"{BOOK_DIR}/{name}.md", encoding="utf-8") as f:
            text = f.read()
        if i == 0:
            # start at the Note on Names (skip title/copyright already built)
            idx = text.find("## A Note on Names")
            text = text[idx:]
        chunks.append(text)
    return "\n\n".join(chunks)

content = read_content()
lines = content.split("\n")

buf = []  # pending list handling
i = 0
n = len(lines)
while i < n:
    line = lines[i].rstrip("\n")
    stripped = line.strip()

    if stripped == "---" or stripped == "":
        i += 1; continue

    # Part headings  (# X)  but not the book title
    if re.match(r'^# ', line):
        title = stripped[2:].strip()
        new_page()
        toc_h(1, title)
        i += 1; continue

    # Chapter / section headings (## X)
    if re.match(r'^## ', line):
        title = stripped[3:].strip()
        new_page()
        toc_h(2, title)
        i += 1; continue

    # Subheadings (### X)
    if re.match(r'^### ', line):
        title = stripped[4:].strip()
        p = doc.add_paragraph(style='SubHead'); add_inline(p, title)
        i += 1; continue

    # Blockquote
    if stripped.startswith(">"):
        qtext = stripped.lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.3)
        p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
        add_inline(p, qtext)
        for r in p.runs: r.italic = True
        i += 1; continue

    # Bullet list
    if stripped.startswith("- "):
        p = doc.add_paragraph(style='List Bullet')
        add_inline(p, stripped[2:].strip())
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1; continue

    # Numbered list
    m = re.match(r'^\d+\.\s+(.*)', stripped)
    if m:
        p = doc.add_paragraph(style='List Number')
        add_inline(p, m.group(1))
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1; continue

    # Normal body paragraph
    p = doc.add_paragraph(style='Body')
    add_inline(p, stripped)
    i += 1

out = f"{BOOK_DIR}/KDP-interior-6x9.docx"
doc.save(out)
print("saved", out)
