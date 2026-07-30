#!/usr/bin/env python3
"""Build the companion '100 Questions About Mormonism' KDP interior (6x9) DOCX."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BOOK="/home/user/sophie-ui/book/companion/es"
PARTS=["01-front-and-intro","s01","s02","s03","s04","s05","s06","s07","s08","s09","s10","99-backmatter"]

def add_field(p,instr):
    r=p.add_run()
    f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text=instr
    f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'separate')
    t=OxmlElement('w:t'); t.text=""
    f3=OxmlElement('w:fldChar'); f3.set(qn('w:fldCharType'),'end')
    for e in (f1,it,f2,t,f3): r._r.append(e)

def add_inline(p,text):
    for tok in re.split(r'(\*\*.*?\*\*|\*.*?\*)',text):
        if not tok: continue
        if tok.startswith('**') and tok.endswith('**'): p.add_run(tok[2:-2]).bold=True
        elif tok.startswith('*') and tok.endswith('*'): p.add_run(tok[1:-1]).italic=True
        else: p.add_run(tok)

doc=Document()
n=doc.styles['Normal']; n.font.name='Georgia'; n.font.size=Pt(11)
pf=n.paragraph_format; pf.line_spacing_rule=WD_LINE_SPACING.MULTIPLE; pf.line_spacing=1.15; pf.space_after=Pt(0); pf.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
body=doc.styles.add_style('Body',1); body.base_style=doc.styles['Normal']; body.font.name='Georgia'; body.font.size=Pt(11)
body.paragraph_format.first_line_indent=Inches(0.22); body.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY; body.paragraph_format.line_spacing=1.15
for name,size,before,after in [('PartHead',22,34,16),('ChapHead',17,28,14)]:
    st=doc.styles.add_style(name,1); st.base_style=doc.styles['Normal']; st.font.name='Georgia'; st.font.size=Pt(size); st.font.bold=True
    st.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER; st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    st.font.color.rgb=RGBColor(0x1a,0x1a,0x1a)
qst=doc.styles.add_style('Qhead',1); qst.base_style=doc.styles['Normal']; qst.font.name='Georgia'; qst.font.size=Pt(12.5); qst.font.bold=True
qst.paragraph_format.space_before=Pt(13); qst.paragraph_format.space_after=Pt(3); qst.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT; qst.paragraph_format.keep_with_next=True

sec=doc.sections[0]; sec.page_width=Inches(6); sec.page_height=Inches(9)
sec.top_margin=Inches(0.75); sec.bottom_margin=Inches(0.75); sec.left_margin=Inches(0.875); sec.right_margin=Inches(0.5); sec.gutter=Inches(0.13)
doc.settings.element.append(OxmlElement('w:mirrorMargins'))
fp=sec.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER; add_field(fp,"PAGE")

def new_page(): doc.add_page_break()
def toc_h(level,text):
    p=doc.add_paragraph(style=('ChapHead' if level==2 else 'PartHead')); add_inline(p,text)
    ol=OxmlElement('w:outlineLvl'); ol.set(qn('w:val'),str(level-1)); p._p.get_or_add_pPr().append(ol)

# TITLE PAGE
for _ in range(3): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("100"); r.bold=True; r.font.size=Pt(66)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("PREGUNTAS SOBRE EL MORMONISMO"); r.bold=True; r.font.size=Pt(20)
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Respuestas claras y honestas sobre los Santos de los\nÚltimos Días: sus creencias, su historia y su vida cotidiana"); r.italic=True; r.font.size=Pt(12.5)
for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("SOPHIE E T"); r.font.size=Pt(15)
for _ in range(4): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Una guía educativa independiente"); r.italic=True; r.font.size=Pt(11)
new_page()

# COPYRIGHT
cp=[
 "Copyright © 2026. Todos los derechos reservados.","",
 "Ninguna parte de este libro puede reproducirse, distribuirse ni transmitirse de ninguna forma ni por ningún medio sin el permiso previo por escrito del editor, salvo en el caso de breves citas incluidas en reseñas críticas y ciertos otros usos no comerciales permitidos por la ley de derechos de autor.","",
 "Nota sobre la independencia. Esta es una obra independiente de divulgación y periodismo. No está publicada, patrocinada, autorizada ni respaldada por La Iglesia de Jesucristo de los Santos de los Últimos Días, ninguna de sus organizaciones afiliadas, ni ninguna otra iglesia, confesión o grupo mencionado en estas páginas. Todas las marcas y nombres son propiedad de sus respectivos titulares y se usan aquí solo con fines de identificación y comentario.","",
 "Nota sobre los nombres. Hoy los miembros prefieren que se les llame Santos de los Últimos Días en vez de mormones. Este libro usa Santo de los Últimos Días como forma respetuosa por defecto y mormón donde ayude a la claridad o refleje el uso común. No se pretende faltar el respeto con ninguno de los dos términos.","",
 "Primera edición — 2026",
]
for _ in range(2): doc.add_paragraph()
for line in cp:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.line_spacing=1.15
    if line:
        m=re.match(r'^(Nota sobre [^.]+\.)(.*)',line)
        if m: p.add_run(m.group(1)).bold=True; p.add_run(m.group(2))
        else: p.add_run(line)
    for rr in p.runs: rr.font.size=Pt(9.5)
new_page()

# TOC
p=doc.add_paragraph(style='ChapHead'); p.add_run("Contenido")
p2=doc.add_paragraph(); add_field(p2,'TOC \\o "1-2" \\h \\z \\u')
doc.settings.element.append(OxmlElement('w:updateFields')); doc.settings.element[-1].set(qn('w:val'),'true')
new_page()

# CONTENT
chunks=[]
for i,name in enumerate(PARTS):
    txt=open(f"{BOOK}/{name}.md",encoding="utf-8").read()
    if i==0:
        idx=txt.find("# Cómo usar este libro")
        if idx>=0: txt=txt[idx:]
    chunks.append(txt)
lines="\n\n".join(chunks).split("\n")
i=0
while i<len(lines):
    line=lines[i].rstrip(); s=line.strip()
    if s=="---" or s=="": i+=1; continue
    if re.match(r'^# ',line): new_page(); toc_h(1,s[2:].strip()); i+=1; continue
    if re.match(r'^## ',line): new_page(); toc_h(2,s[3:].strip()); i+=1; continue
    if re.match(r'^### ',line):
        p=doc.add_paragraph(style='Qhead'); add_inline(p,s[4:].strip()); i+=1; continue
    if s.startswith("- "):
        p=doc.add_paragraph(style='List Bullet'); add_inline(p,s[2:].strip()); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT; i+=1; continue
    m=re.match(r'^\d+\.\s+(.*)',s)
    if m and not re.match(r'^\d+\.\s',s) is None and len(s)<200 and s.endswith('?'):
        # numbered question line that slipped through as body (safety) -> Qhead
        p=doc.add_paragraph(style='Qhead'); add_inline(p,s); i+=1; continue
    p=doc.add_paragraph(style='Body'); add_inline(p,s); i+=1

doc.save(f"{BOOK}/KDP-interior-6x9.docx"); print("saved companion/KDP-interior-6x9.docx")
